"""Context ingest — drag/drop PDFs, DOCX, text, or ChatGPT exports.

Flow (now async to dodge Railway/Cloudflare's ~100s HTTP timeout):

  1. HTTP handler parses every uploaded file synchronously (cheap — just
     PDF/DOCX text extraction, no network).
  2. For each file we insert a `documents` row in `processing` state and
     return immediately. The frontend already subscribes to that table via
     Supabase Realtime, so rows appear as 'processing' instantly and flip
     to 'complete' (or 'failed') when the background task finishes.
  3. The background task batches embeddings (~64 chunks per OpenRouter call),
     upserts them to Qdrant in one go per batch, and finally extracts identity
     facts to Mem0.
"""

from __future__ import annotations

import asyncio
import io
import json
import zipfile
from typing import Any

import structlog
from fastapi import APIRouter, BackgroundTasks, File, HTTPException, UploadFile
from pydantic import BaseModel

from ..auth import CurrentUserDep
from ..audit import write as audit_write
from ..db import supabase_admin
from ..memory.mem0 import JaiMem0
from ..memory.neo4j_client import JaiNeo4j
from ..memory.qdrant import JaiQdrant

log = structlog.get_logger()
router = APIRouter()


# --- limits --------------------------------------------------------------

MAX_FILE_MB = 25
MAX_TOTAL_CHUNKS = 5000          # per-request safety cap (was 2000 — now async)
CHUNK_CHARS = 1200
CHUNK_OVERLAP = 200
MAX_GPT_CONVOS = 1000
EMBED_BATCH = 64                 # chunks per embedding API call


# --- text extraction -----------------------------------------------------


def _extract_pdf(blob: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(500, "pypdf not installed on the server")
    try:
        reader = PdfReader(io.BytesIO(blob))
    except Exception as e:
        log.warning("pdf.read_failed", error=str(e))
        return ""
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
            if txt.strip():
                pages.append(f"[page {i + 1}]\n{txt.strip()}")
        except Exception as e:
            log.warning("pdf.page_failed", page=i + 1, error=str(e))
    return "\n\n".join(pages)


def _extract_docx(blob: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(500, "python-docx not installed on the server")
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as e:
        log.warning("docx.read_failed", error=str(e))
        return ""
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_chatgpt_conversations(data: Any) -> list[dict[str, Any]]:
    if not isinstance(data, list):
        return []
    out: list[dict[str, Any]] = []
    for convo in data[:MAX_GPT_CONVOS]:
        if not isinstance(convo, dict):
            continue
        title = (convo.get("title") or "").strip() or "Untitled"
        created = convo.get("create_time")
        mapping = convo.get("mapping") or {}
        user_msgs: list[str] = []
        for node in mapping.values():
            if not isinstance(node, dict):
                continue
            msg = node.get("message") or {}
            author = (msg.get("author") or {}).get("role")
            if author != "user":
                continue
            parts = (msg.get("content") or {}).get("parts") or []
            for p in parts:
                if isinstance(p, str) and p.strip():
                    user_msgs.append(p.strip())
        if not user_msgs:
            continue
        out.append({"title": title, "created": created, "body": "\n\n".join(user_msgs)})
    return out


def _extract_from_zip(blob: bytes) -> tuple[str, list[dict[str, Any]]]:
    plain: list[str] = []
    convos: list[dict[str, Any]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as z:
            for info in z.infolist():
                if info.is_dir() or info.file_size > MAX_FILE_MB * 1024 * 1024:
                    continue
                name = info.filename.lower()
                if name.endswith("conversations.json"):
                    try:
                        data = json.loads(z.read(info.filename).decode("utf-8"))
                        convos.extend(_extract_chatgpt_conversations(data))
                    except Exception as e:
                        log.warning("zip.gpt_parse_failed", file=info.filename, error=str(e))
                elif name.endswith((".txt", ".md")):
                    try:
                        plain.append(z.read(info.filename).decode("utf-8", errors="ignore"))
                    except Exception:
                        pass
                elif name.endswith(".pdf"):
                    try:
                        plain.append(_extract_pdf(z.read(info.filename)))
                    except Exception:
                        pass
                elif name.endswith(".docx"):
                    try:
                        plain.append(_extract_docx(z.read(info.filename)))
                    except Exception:
                        pass
    except zipfile.BadZipFile:
        return "", []
    return "\n\n".join(p for p in plain if p.strip()), convos


# --- chunking ------------------------------------------------------------


def _chunk(text: str, size: int = CHUNK_CHARS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = text.strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap
    return chunks


# --- fact extraction -----------------------------------------------------


async def _extract_entities(text: str) -> dict[str, Any]:
    """Pull entities + relationships out of a blob (best-effort).

    Returns shape:
      {
        "people":   [{"name": "...", "role": "..."}, ...],
        "companies":[{"name": "...", "note": "..."}, ...],
        "projects": [{"name": "...", "note": "..."}, ...],
        "beliefs":  [{"text": "..."}, ...],
        "topics":   [{"name": "..."}, ...],
        "relations":[{"src": "...", "rel": "WORKS_WITH", "dst": "..."}, ...],
      }
    """
    if not text.strip():
        return {}
    try:
        from ..config import get_settings
        from ..models.openrouter import openrouter_chat

        s = get_settings()
        llm = openrouter_chat(
            model=s.jai_model_orchestrator,
            settings=s,
            temperature=0.0,
            streaming=False,
        )
        prompt = (
            "Extract entities and relationships from the text. Return STRICT JSON "
            'matching this shape (no prose, no markdown fences):\n'
            "{\n"
            '  "people": [{"name": str, "role": str|null}],\n'
            '  "companies": [{"name": str, "note": str|null}],\n'
            '  "projects": [{"name": str, "note": str|null}],\n'
            '  "beliefs": [{"text": str}],\n'
            '  "topics": [{"name": str}],\n'
            '  "relations": [{"src": str, "rel": str, "dst": str}]\n'
            "}\n"
            "Rules: include only entities clearly mentioned. Use SCREAMING_SNAKE_CASE "
            'for relation types (e.g. "WORKS_WITH", "FOUNDED", "BELIEVES").\n'
            "If a category is empty, return [].\n\n"
            "TEXT:\n"
            f"{text[:12000]}"
        )
        resp = await llm.ainvoke(prompt)
        raw = str(resp.content or "").strip()

        # Tolerate accidental code fences.
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:].lstrip()
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            # Try to slice out the first {...} block.
            start = raw.find("{")
            end = raw.rfind("}")
            if start >= 0 and end > start:
                try:
                    data = json.loads(raw[start : end + 1])
                except json.JSONDecodeError:
                    return {}
            else:
                return {}
        return data if isinstance(data, dict) else {}
    except Exception as e:
        log.warning("ingest.entity_extract_failed", error=str(e))
        return {}


def _slug(s: str) -> str:
    """Stable slug for entity ids derived from names."""
    return "".join(c if c.isalnum() else "_" for c in s.lower()).strip("_")[:80] or "x"


async def _write_to_neo4j(user_id: str, ents: dict[str, Any]) -> tuple[int, int]:
    """Upsert extracted entities and relations. Returns (nodes, edges) written."""
    if not ents:
        return 0, 0
    try:
        n4 = JaiNeo4j()
    except Exception as e:
        log.warning("neo4j.init_failed", error=str(e))
        return 0, 0

    nodes = 0
    edges = 0
    name_to_label: dict[str, str] = {}

    async def upsert_collection(items: list[dict[str, Any]] | None, label: str) -> int:
        if not isinstance(items, list):
            return 0
        count = 0
        for it in items:
            name = (it or {}).get("name") if label != "Belief" else (it or {}).get("text")
            if not name or not isinstance(name, str):
                continue
            entity_id = f"{label.lower()}_{_slug(name)}"
            props: dict[str, Any] = {"name": name}
            # Merge any extra optional fields (role/note) without clobbering.
            for k, v in (it or {}).items():
                if k != "name" and v is not None:
                    props[k] = v
            try:
                await n4.upsert_entity(
                    user_id=user_id, label=label, entity_id=entity_id, properties=props
                )
                name_to_label[name.lower()] = (label, entity_id)  # type: ignore[assignment]
                count += 1
            except Exception as e:
                log.warning("neo4j.upsert_node_failed", label=label, error=str(e))
        return count

    try:
        nodes += await upsert_collection(ents.get("people"), "Person")
        nodes += await upsert_collection(ents.get("companies"), "Company")
        nodes += await upsert_collection(ents.get("projects"), "Project")
        nodes += await upsert_collection(ents.get("beliefs"), "Belief")
        nodes += await upsert_collection(ents.get("topics"), "Topic")

        for rel in ents.get("relations") or []:
            src_name = (rel or {}).get("src")
            dst_name = (rel or {}).get("dst")
            rel_type = (rel or {}).get("rel")
            if not (
                isinstance(src_name, str)
                and isinstance(dst_name, str)
                and isinstance(rel_type, str)
            ):
                continue
            # Sanitize rel type — Cypher relationship types must be identifiers.
            rel_safe = "".join(c if (c.isalnum() or c == "_") else "_" for c in rel_type).upper()
            if not rel_safe or rel_safe[0].isdigit():
                rel_safe = f"R_{rel_safe}"
            src = name_to_label.get(src_name.lower())
            dst = name_to_label.get(dst_name.lower())
            if not (src and dst):
                continue
            try:
                await n4.upsert_rel(
                    user_id=user_id,
                    src_label=src[0],
                    src_id=src[1],
                    rel_type=rel_safe,
                    dst_label=dst[0],
                    dst_id=dst[1],
                )
                edges += 1
            except Exception as e:
                log.warning("neo4j.upsert_rel_failed", error=str(e))
    finally:
        try:
            await n4.close()
        except Exception:
            pass

    return nodes, edges


async def _extract_facts(text: str, max_facts: int = 12) -> list[str]:
    if not text.strip():
        return []
    try:
        from ..config import get_settings
        from ..models.openrouter import openrouter_chat

        s = get_settings()
        llm = openrouter_chat(model=s.jai_model_orchestrator, settings=s, temperature=0.1)
        prompt = (
            "Extract up to "
            f"{max_facts} DURABLE, IDENTITY-SHAPED facts about the human from the text. "
            "Keep ONLY facts that would matter to a personal AI six months from now: "
            "core identity, role, business/company, recurring beliefs, deep "
            "preferences, named relationships, ongoing goals, values, signature "
            "ways of thinking. \n\n"
            "REJECT and skip: one-off tasks, today's todos, casual questions, "
            "random metrics they mentioned once (e.g. 'ShoutOut grew 300%' is a "
            "stat, not an identity fact unless framed as a pattern), administrative "
            "trivia, anything that reads like a transient note. \n\n"
            "Return one fact per line, no numbering, no preamble. If nothing meets "
            "the bar, return ONLY the literal string NONE on its own line.\n\n"
            "TEXT:\n"
            f"{text[:8000]}"
        )
        resp = await llm.ainvoke(prompt)
        raw = str(resp.content or "")
        if raw.strip().upper().startswith("NONE"):
            return []
        facts = [
            line.strip().lstrip("-•* ").strip()
            for line in raw.splitlines()
            if line.strip() and len(line.strip()) > 5 and line.strip().upper() != "NONE"
        ]
        return facts[:max_facts]
    except Exception as e:
        log.warning("ingest.fact_extract_failed", error=str(e))
        return []


# --- API models ----------------------------------------------------------


class IngestStub(BaseModel):
    """What we return synchronously — one entry per accepted file."""
    document_id: str
    filename: str
    status: str           # 'processing' | 'failed'
    error: str | None = None


class IngestResponse(BaseModel):
    accepted: list[IngestStub]
    skipped: list[str] = []


class DocumentRow(BaseModel):
    id: str
    filename: str
    size_bytes: int | None = None
    content_type: str | None = None
    kind: str
    status: str = "complete"
    chunks_count: int
    conversations_count: int
    facts_count: int
    metadata: dict[str, Any] | None = None
    error: str | None = None
    created_at: str


class SearchHit(BaseModel):
    text: str
    score: float
    source: str | None = None
    filename: str | None = None
    title: str | None = None


class SearchIn(BaseModel):
    q: str
    limit: int | None = None


# --- background worker ---------------------------------------------------


async def _process_one(
    *,
    user_id: str,
    doc_id: str,
    filename: str,
    plain_text: str,
    gpt_convos: list[dict[str, Any]],
) -> None:
    """Embed → Qdrant → Mem0. Updates the documents row with final counts."""
    sb = supabase_admin()
    qd = JaiQdrant()
    mem = JaiMem0()

    chunks_count = 0
    conv_count = 0
    facts_count = 0

    try:
        await qd.ensure_collection()

        # --- ChatGPT conversations -----------------------------------
        if gpt_convos:
            items: list[dict[str, Any]] = []
            for c in gpt_convos:
                for chunk in _chunk(c["body"]):
                    items.append(
                        {
                            "text": chunk,
                            "source": "chatgpt_export",
                            "metadata": {
                                "title": c["title"],
                                "created": c.get("created"),
                                "kind": "chat_history",
                                "filename": filename,
                            },
                        }
                    )
                    if len(items) >= MAX_TOTAL_CHUNKS:
                        break
                conv_count += 1
                if len(items) >= MAX_TOTAL_CHUNKS:
                    break
            if items:
                chunks_count += await qd.add_batch(
                    user_id=user_id, items=items, batch_size=EMBED_BATCH
                )

        # --- plain text ----------------------------------------------
        if plain_text.strip():
            items = [
                {
                    "text": ch,
                    "source": f"upload:{filename}",
                    "metadata": {"kind": "document", "filename": filename},
                }
                for ch in _chunk(plain_text)
            ][: max(0, MAX_TOTAL_CHUNKS - chunks_count)]
            if items:
                chunks_count += await qd.add_batch(
                    user_id=user_id, items=items, batch_size=EMBED_BATCH
                )

            # Fact extraction once per file
            facts = await _extract_facts(plain_text)
            if facts:
                try:
                    msgs = [{"role": "user", "content": f"Fact about me: {x}"} for x in facts]
                    await mem.add(user_id, msgs, metadata={"source": f"upload:{filename}"})
                    facts_count = len(facts)
                except Exception as e:
                    log.warning("ingest.mem0_add_failed", error=str(e))

        # --- entity extraction → Neo4j -------------------------------
        # Run over BOTH plain text and ChatGPT export bodies so the
        # identity graph grows from chat-history uploads too. Cap the
        # input at 12k chars (matches `_extract_entities`).
        graph_source: list[str] = []
        if plain_text.strip():
            graph_source.append(plain_text)
        if gpt_convos:
            for c in gpt_convos[:50]:  # cap to keep the prompt sane
                body = c.get("body") or ""
                if body:
                    graph_source.append(body)
        combined = "\n\n".join(graph_source)[:12000]
        if combined.strip():
            try:
                ents = await _extract_entities(combined)
                if ents:
                    nodes_w, edges_w = await _write_to_neo4j(user_id, ents)
                    log.info(
                        "ingest.graph_populated",
                        filename=filename,
                        nodes=nodes_w,
                        edges=edges_w,
                    )
            except Exception as e:
                log.warning("ingest.graph_failed", error=str(e))

        sb.table("documents").update(
            {
                "status": "complete",
                "chunks_count": chunks_count,
                "conversations_count": conv_count,
                "facts_count": facts_count,
            }
        ).eq("id", doc_id).eq("user_id", user_id).execute()

        # Mark onboarded after first successful ingest
        try:
            sb.table("users").update({"metadata": {"onboarded": True}}).eq(
                "id", user_id
            ).execute()
        except Exception:
            pass

        await audit_write(
            user_id=user_id,
            actor="context",
            action="context.ingest.complete",
            payload={
                "document_id": doc_id,
                "filename": filename,
                "chunks": chunks_count,
                "conversations": conv_count,
                "facts": facts_count,
            },
        )
    except Exception as e:
        log.exception("ingest.background_failed", document_id=doc_id, filename=filename)
        try:
            sb.table("documents").update(
                {
                    "status": "failed",
                    "error": str(e)[:500],
                    "chunks_count": chunks_count,
                    "conversations_count": conv_count,
                    "facts_count": facts_count,
                }
            ).eq("id", doc_id).eq("user_id", user_id).execute()
        except Exception:
            pass


# --- routes --------------------------------------------------------------


@router.post("/ingest", response_model=IngestResponse)
async def ingest(
    user: CurrentUserDep,
    background: BackgroundTasks,
    files: list[UploadFile] = File(...),
) -> IngestResponse:
    if not files:
        raise HTTPException(400, "no files uploaded")

    sb = supabase_admin()
    accepted: list[IngestStub] = []
    skipped: list[str] = []

    for f in files:
        name = f.filename or "unnamed"
        ctype = (f.content_type or "").lower()
        blob = await f.read()
        if len(blob) > MAX_FILE_MB * 1024 * 1024:
            skipped.append(f"{name} (>{MAX_FILE_MB}MB)")
            continue

        plain_text = ""
        gpt_convos: list[dict[str, Any]] = []

        lower = name.lower()
        try:
            if ctype == "application/pdf" or lower.endswith(".pdf"):
                plain_text = _extract_pdf(blob)
            elif (
                ctype
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or lower.endswith(".docx")
            ):
                plain_text = _extract_docx(blob)
            elif ctype == "application/zip" or lower.endswith(".zip"):
                plain_text, gpt_convos = _extract_from_zip(blob)
            elif ctype == "application/json" or lower.endswith(".json"):
                data = json.loads(blob.decode("utf-8"))
                gpt_convos = _extract_chatgpt_conversations(data)
                if not gpt_convos:
                    plain_text = blob.decode("utf-8", errors="ignore")
            elif ctype.startswith("text/") or lower.endswith((".txt", ".md")):
                plain_text = blob.decode("utf-8", errors="ignore")
            else:
                skipped.append(f"{name} (unsupported type {ctype or '?'})")
                continue
        except Exception as e:
            skipped.append(f"{name} ({type(e).__name__}: {e})")
            continue

        if not plain_text.strip() and not gpt_convos:
            skipped.append(f"{name} (no extractable text)")
            continue

        kind = "chatgpt_export" if gpt_convos else "document"

        # Insert the row immediately so the user sees it in the UI right away.
        try:
            res = (
                sb.table("documents")
                .insert(
                    {
                        "user_id": user.user_id,
                        "filename": name,
                        "size_bytes": len(blob),
                        "content_type": ctype or None,
                        "kind": kind,
                        "status": "processing",
                        "chunks_count": 0,
                        "conversations_count": 0,
                        "facts_count": 0,
                    }
                )
                .execute()
            )
            doc_id = res.data[0]["id"]
        except Exception as e:
            skipped.append(f"{name} (db insert: {e})")
            continue

        # Hand the heavy work to FastAPI's background scheduler.
        background.add_task(
            _process_one,
            user_id=user.user_id,
            doc_id=doc_id,
            filename=name,
            plain_text=plain_text,
            gpt_convos=gpt_convos,
        )

        accepted.append(
            IngestStub(document_id=doc_id, filename=name, status="processing")
        )

    return IngestResponse(accepted=accepted, skipped=skipped)


@router.get("/documents", response_model=list[DocumentRow])
async def list_documents(user: CurrentUserDep, limit: int = 200) -> list[dict[str, Any]]:
    sb = supabase_admin()
    res = (
        sb.table("documents")
        .select("*")
        .eq("user_id", user.user_id)
        .order("created_at", desc=True)
        .limit(min(max(limit, 1), 500))
        .execute()
    )
    return res.data or []


@router.delete("/documents/{doc_id}")
async def delete_document(user: CurrentUserDep, doc_id: str) -> dict[str, Any]:
    sb = supabase_admin()
    sb.table("documents").delete().eq("user_id", user.user_id).eq("id", doc_id).execute()
    return {"ok": True}


@router.post("/documents/sweep_stuck")
async def sweep_stuck_documents(user: CurrentUserDep, older_than_min: int = 10) -> dict[str, Any]:
    """Mark any document stuck in 'processing' for >N min as failed.

    Background ingest tasks die silently when the FastAPI worker restarts
    (Railway redeploys, OOM, etc.), leaving rows in 'processing' forever.
    Run this on every doc-list fetch from the UI so the badge reflects
    reality and the user can retry.
    """
    from datetime import datetime, timedelta, timezone
    cutoff = (datetime.now(timezone.utc) - timedelta(minutes=older_than_min)).isoformat()
    sb = supabase_admin()
    res = (
        sb.table("documents")
        .update({"status": "failed", "error": "ingest worker timed out — click retry"})
        .eq("user_id", user.user_id)
        .eq("status", "processing")
        .lt("created_at", cutoff)
        .execute()
    )
    return {"ok": True, "swept": len(res.data or [])}


# --- graph CRUD ----------------------------------------------------------


class GraphNodePatch(BaseModel):
    name: str


@router.delete("/graph/node/{node_id}")
async def delete_graph_node(user: CurrentUserDep, node_id: str) -> dict[str, Any]:
    """Detach-delete one node from the user's identity graph."""
    try:
        n4 = JaiNeo4j()
    except Exception as e:
        raise HTTPException(503, f"graph unavailable: {e}") from e
    try:
        deleted = await n4.delete_entity(user_id=user.user_id, node_id=node_id)
        if deleted == 0:
            raise HTTPException(404, "node not found")
        await audit_write(
            user_id=user.user_id,
            actor="context",
            action="context.graph.node.delete",
            payload={"node_id": node_id},
        )
        return {"ok": True, "deleted": deleted}
    finally:
        try:
            await n4.close()
        except Exception:
            pass


@router.patch("/graph/node/{node_id}")
async def patch_graph_node(
    user: CurrentUserDep, node_id: str, body: GraphNodePatch
) -> dict[str, Any]:
    """Rename a node (the only safe in-place edit for now)."""
    name = (body.name or "").strip()
    if not name:
        raise HTTPException(400, "name required")
    try:
        n4 = JaiNeo4j()
    except Exception as e:
        raise HTTPException(503, f"graph unavailable: {e}") from e
    try:
        ok = await n4.update_entity_name(
            user_id=user.user_id, node_id=node_id, name=name
        )
        if not ok:
            raise HTTPException(404, "node not found")
        return {"ok": True, "name": name}
    finally:
        try:
            await n4.close()
        except Exception:
            pass


class ContextForget(BaseModel):
    query: str
    scopes: list[str] | None = None   # ["qdrant","graph","mem0"] — default all


@router.post("/forget")
async def context_forget(user: CurrentUserDep, body: ContextForget) -> dict[str, Any]:
    """One-shot 'remove this from my context' that nukes everything matching.

    Used by the chat fast-path when the user says e.g. "remove ShoutOut's 300%
    growth from context". Hits all three memory tiers and reports what went.
    """
    q = (body.query or "").strip()
    if not q:
        raise HTTPException(400, "query required")
    scopes = set(body.scopes or ["qdrant", "graph", "mem0"])
    out: dict[str, Any] = {"query": q, "deleted": {}}

    if "qdrant" in scopes:
        qd = JaiQdrant()
        try:
            await qd.ensure_collection()
            n = await qd.delete_by_query(user_id=user.user_id, query=q, top_k=25)
            out["deleted"]["qdrant"] = n
        finally:
            try:
                await qd.close()
            except Exception:
                pass

    if "graph" in scopes:
        try:
            n4 = JaiNeo4j()
            try:
                # Try to match by name fragment — keeps the API simple.
                nodes = await n4.delete_by_name_fragment(
                    user_id=user.user_id, fragment=q
                )
                out["deleted"]["graph"] = [n["name"] for n in nodes]
            finally:
                try:
                    await n4.close()
                except Exception:
                    pass
        except Exception as e:
            log.warning("forget.graph_failed", error=str(e))
            out["deleted"]["graph"] = []

    if "mem0" in scopes:
        try:
            mem = JaiMem0()
            out["deleted"]["mem0"] = await mem.delete_about(user.user_id, q)
        except Exception as e:
            log.warning("forget.mem0_failed", error=str(e))
            out["deleted"]["mem0"] = 0

    await audit_write(
        user_id=user.user_id,
        actor="context",
        action="context.forget",
        payload=out,
    )
    return out


@router.post("/graph/rebuild")
async def graph_rebuild(user: CurrentUserDep) -> dict[str, Any]:
    """Re-run entity extraction over every previously-ingested chunk in Qdrant
    and seed the Neo4j graph. Idempotent — MERGE-based upserts in Cypher mean
    repeated runs converge instead of duplicating.

    Runs synchronously so the UI can report real counts. We cap the number
    of documents we process per call to keep it under Railway's ~100s
    response budget; the user can click again to keep going.
    """
    qd = JaiQdrant()
    docs_seen = 0
    chunks_seen = 0
    total_nodes = 0
    total_edges = 0
    skipped: list[str] = []

    try:
        await qd.ensure_collection()

        from qdrant_client.http.models import FieldCondition, Filter, MatchValue
        flt = Filter(
            must=[FieldCondition(key="user_id", match=MatchValue(value=user.user_id))]
        )
        buckets: dict[str, list[str]] = {}
        offset = None
        # 50 pages of 200 = 10k chunks ceiling — way more than we'd ever
        # have in practice for a single user.
        for _ in range(50):
            pts, offset = await qd._client.scroll(  # noqa: SLF001
                collection_name=qd._collection,  # noqa: SLF001
                scroll_filter=flt,
                limit=200,
                offset=offset,
                with_payload=True,
                with_vectors=False,
            )
            for p in pts:
                payload = p.payload or {}
                src = (
                    payload.get("source")
                    or payload.get("filename")
                    or "untagged"
                )
                text = payload.get("text") or ""
                if text:
                    buckets.setdefault(src, []).append(text)
                    chunks_seen += 1
            if not offset:
                break

        # Cap per-call work so the request fits inside Railway's HTTP budget.
        # Each doc costs ~1 LLM call + a handful of Neo4j writes (~5-15s total).
        MAX_DOCS_PER_CALL = 8
        items = list(buckets.items())[:MAX_DOCS_PER_CALL]

        for src, chunks in items:
            joined = "\n\n".join(chunks)[:12000]
            try:
                ents = await _extract_entities(joined)
            except Exception as e:
                skipped.append(f"{src}: extract failed: {e}")
                continue
            if not ents:
                skipped.append(f"{src}: no entities")
                continue
            try:
                n, e = await _write_to_neo4j(user.user_id, ents)
            except Exception as exc:
                skipped.append(f"{src}: neo4j failed: {exc}")
                continue
            total_nodes += n
            total_edges += e
            docs_seen += 1
            log.info("graph.rebuild.doc", source=src, nodes=n, edges=e)

        try:
            await audit_write(
                user_id=user.user_id,
                actor="context",
                action="context.graph.rebuild",
                payload={
                    "docs_processed": docs_seen,
                    "docs_total": len(buckets),
                    "chunks": chunks_seen,
                    "nodes": total_nodes,
                    "edges": total_edges,
                },
            )
        except Exception:
            pass

        return {
            "ok": True,
            "docs_processed": docs_seen,
            "docs_total": len(buckets),
            "chunks_scanned": chunks_seen,
            "nodes_written": total_nodes,
            "edges_written": total_edges,
            "remaining": max(0, len(buckets) - docs_seen),
            "skipped": skipped,
        }
    except Exception as e:
        log.exception("graph.rebuild.failed")
        raise HTTPException(500, f"rebuild failed: {e}") from e
    finally:
        try:
            await qd.close()
        except Exception:
            pass


@router.get("/graph")
async def graph(user: CurrentUserDep, limit: int = 200) -> dict[str, Any]:
    """Return the user's identity graph (Neo4j Aura) for the Context → Graph view."""
    try:
        n4 = JaiNeo4j()
    except Exception as e:
        # Don't 500 the UI; return empty so the empty-state copy shows.
        log.warning("graph.init_failed", error=str(e))
        return {"nodes": [], "edges": []}
    try:
        return await n4.graph_dump(user.user_id, limit=min(max(limit, 1), 500))
    except Exception as e:
        log.warning("graph.dump_failed", error=str(e))
        return {"nodes": [], "edges": []}
    finally:
        try:
            await n4.close()
        except Exception:
            pass


@router.post("/search", response_model=list[SearchHit])
async def search(user: CurrentUserDep, body: SearchIn) -> list[dict[str, Any]]:
    q = (body.q or "").strip()
    if not q:
        raise HTTPException(400, "missing query")
    qd = JaiQdrant()
    try:
        await qd.ensure_collection()
        hits = await qd.search(user.user_id, q)
    except Exception as e:
        log.warning("context.search_failed", error=str(e))
        raise HTTPException(500, f"search failed: {e}") from e

    limit = min(max(body.limit or 20, 1), 50)
    out: list[dict[str, Any]] = []
    for h in hits[:limit]:
        meta = h.get("metadata") or {}
        out.append(
            {
                "text": h.get("text", ""),
                "score": float(h.get("score", 0.0)),
                "source": meta.get("source"),
                "filename": meta.get("filename"),
                "title": meta.get("title"),
            }
        )
    return out


# Quiet down a lint complaint about an unused import only relevant once we
# extend the worker with explicit concurrency primitives.
_ = asyncio
